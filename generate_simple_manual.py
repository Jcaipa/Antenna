from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_simple_guide():
    doc = Document()

    # Título Llamativo
    title = doc.add_heading('🚀 Guía Rápida: Cómo usar tu Sistema Antenna', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Este sistema vigila tu marca y competencia las 24 horas. Sigue estos pasos para personalizarlo a tu gusto.')

    # 1. ¿Cómo cambiar los temas y marcas?
    doc.add_heading('1. 🎯 Cambiar qué quieres vigilar (Personalización)', level=1)
    doc.add_paragraph('Si quieres que el sistema hable de otros temas, abre estos archivos y busca las listas de palabras clave:')
    
    # Tabla de personalización detallada
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Plataforma'
    hdr_cells[1].text = 'Archivo a Modificar'
    hdr_cells[2].text = '¿Para qué sirve este cambio?'

    items = [
        ("Reddit", "Social_listening/reddit.py", "Cambia los 'Subreddits' para escuchar comunidades distintas (ej: de 'Noticias' a 'Turismo')."),
        ("YouTube", "Social_listening/youtube.py", "Cambia las 'Queries' para buscar videos de otros temas (ej: de 'Seguridad' a 'Hoteles')."),
        ("Google Search", "SEO / AEO/serp_rankings.py", "Cambia las 'Keywords' para ver en qué posición sales en Google con nuevas palabras."),
        ("Competencia", "Competitive Intelligence/competitor_monitor.py", "Cambia las 'URLs'. El sistema te avisará si tus competidores cambian sus precios o textos.")
    ]

    for plat, loc, why in items:
        row_cells = table.add_row().cells
        row_cells[0].text = plat
        row_cells[1].text = loc
        row_cells[2].text = why

    # 2. ¿Cómo activar la Inteligencia Artificial?
    doc.add_heading('2. 🧠 Activar la Inteligencia Artificial (IA)', level=1)
    doc.add_paragraph('Para que el sistema genere reportes estratégicos automáticos, pega tus llaves (API Keys) aquí:')
    
    keys = [
        ("Cerebro de IA (Claude)", "Pega tu llave en narrative_detector.py y ad_analyzer.py. Es la que redacta los reportes."),
        ("Anuncios de Facebook", "Pega tu Token en meta_ads.py. Es necesario para 'espiar' la publicidad activa."),
        ("Autoridad Web", "Ya configuramos tu llave de Open PageRank en authority_check.py.")
    ]

    for k_name, k_loc in keys:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(f"{k_name}: ")
        run.bold = True
        p.add_run(f"Ir a {k_loc}")

    # 3. ¿Dónde están mis reportes?
    doc.add_heading('3. 📊 Consultar Resultados', level=1)
    doc.add_paragraph('Los resultados se guardan automáticamente en archivos de Excel (.csv):')
    
    reports = [
        ("Opiniones y Noticias", "Social_listening/reddit_us_insights.csv"),
        ("Ranking en Google", "SEO / AEO/serp_rankings_audit.csv"),
        ("Análisis de la Competencia", "Competitive Intelligence/competitor_tech_stacks.csv")
    ]

    for r_name, r_path in reports:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{r_name}: ")
        run.bold = True
        p.add_run(f"Archivo: {r_path}")

    # Final
    doc.add_heading('¡Listo! 🏁', level=1)
    doc.add_paragraph('Cualquier cambio que hagas en los archivos se aplicará la próxima vez que ejecutes el programa.')

    # Guardar
    file_name = "Guia_Rapida_Antenna.docx"
    doc.save(file_name)
    return file_name

if __name__ == "__main__":
    name = create_simple_guide()
    print(f"✅ Guía rápida regenerada: {name}")
