from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    doc = Document()

    # Titulo Principal
    title = doc.add_heading('Manual del Sistema de Inteligencia "Antenna"', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Este documento resume los módulos construidos para el ecosistema de inteligencia de Antenna, su funcionamiento y las credenciales pendientes para activar el sistema al 100%.')

    # SECCION 1: MODULOS
    doc.add_heading('1. Módulos Implementados', level=1)
    
    modules = [
        ("Social Listening", "Monitorea Reddit, YouTube y Google News para medir el sentimiento y la percepción internacional de EE.UU."),
        ("SEO / AEO (Search & Answers)", "Analiza el ranking en Google y la visibilidad de marca en motores de respuesta de IA (ChatGPT, Claude, Perplexity)."),
        ("Paid Signals", "Espía la publicidad de la competencia en Meta (FB/IG) y Google Ads, analizando sus estrategias creativas."),
        ("Trends Engine", "Detecta narrativas emergentes cruzando tendencias de búsqueda en Google con noticias de vanguardia (Hacker News)."),
        ("Competitive Intelligence", "Vigila cambios en las webs de competidores y analiza su autoridad de dominio y stack tecnológico.")
    ]

    for mod_title, mod_desc in modules:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{mod_title}: ")
        run.bold = True
        p.add_run(mod_desc)

    # SECCION 2: CREDENCIALES
    doc.add_heading('2. Credenciales Pendientes (Para Activar todo)', level=1)
    
    doc.add_paragraph('Aunque el sistema ya genera reportes, algunos módulos requieren "llaves" oficiales para evitar bloqueos y activar la IA:')

    keys = [
        ("Reddit (Evitar Error 429)", "Necesitamos Client ID y Secret. Ve a reddit.com/prefs/apps, crea una 'app' tipo 'script', y pon http://localhost:8080 en redirect URI."),
        ("Meta Ad Library", "Necesitamos un User Access Token. Ve a developers.facebook.com/tools/explorer y genera un token con el permiso 'ads_read'."),
        ("Anthropic (Claude IA)", "Necesitamos una API Key. Esta es la más importante ya que es el 'cerebro' que analiza todos los datos recolectados. Ve a console.anthropic.com."),
        ("OpenAI / Perplexity", "Opcionales para el módulo de AEO si deseas comparar respuestas de otros modelos.")
    ]

    for key_title, key_desc in keys:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(f"{key_title}: ")
        run.bold = True
        p.add_run(key_desc)

    # SECCION 3: EJECUCION
    doc.add_heading('3. Cómo Ejecutar y Enviar', level=1)
    doc.add_paragraph('Cada módulo tiene un "Dashboard" o script maestro que corre todas las herramientas de esa carpeta:')
    
    execs = [
        ("Social Listening", "Corre cada script en Social_listening/"),
        ("SEO / AEO", "Ejecuta python3 gsc_extract.py despues de poner tu URL."),
        ("Paid Signals", "Ejecuta el ad_analyzer.py una vez tengas la llave de Claude."),
        ("Competitive Intelligence", "Corre Competitive Intelligence/comp_intel_dashboard.py")
    ]

    for ex_title, ex_desc in execs:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{ex_title}: ")
        run.bold = True
        p.add_run(ex_desc)

    # SECCION 4: PERSONALIZACION
    doc.add_heading('4. Cómo Personalizar Temas y Competidores', level=1)
    doc.add_paragraph('Si quieres cambiar lo que el sistema vigila (temas, palabras o marcas), busca estas líneas al principio de cada archivo:')
    
    customs = [
        ("Reddit", "Social_listening/reddit.py -> SUBREDDITS = [...]", "¿Qué cambia? Las comunidades donde buscamos opiniones. Cambialo para pasar de buscar 'Noticias' a buscar 'Viajes'."),
        ("YouTube", "Social_listening/youtube.py -> SEARCH_QUERIES = [...]", "¿Qué cambia? Lo que buscamos en YouTube. Cambia 'Seguridad' por 'Turismo' para ver videos de ese tema."),
        ("SEO Rankings", "SEO / AEO/serp_rankings.py -> KEYWORDS = [...]", "¿Qué cambia? Las palabras con las que quieres salir en Google. Úsalo para auditar nuevas palabras."),
        ("Competidores", "Competitive Intelligence/competitor_monitor.py -> COMPETITORS = {...}", "¿Qué cambia? Las páginas de la competencia que espiamos. El sistema te avisará si cambian algo (precios, textos, etc) en esos links.")
    ]

    for c_title, c_loc, c_why in customs:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{c_title}: ")
        run.bold = True
        p.add_run(f"Modifica el archivo en {c_loc}")
        p.add_run(f"\n{c_why}").italic = True

    # Guardar
    file_name = "Guia_Sistema_Antenna.docx"
    doc.save(file_name)
    return file_name

if __name__ == "__main__":
    name = create_manual()
    print(f"✅ Documento generado: {name}")
